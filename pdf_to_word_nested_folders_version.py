import os
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configuration
PDF_FOLDER = 'backup'  # main folder containing patient folders
OUTPUT_FOLDER = 'Subfolder'
LOGO_PATH = 'logo.png'
DESCRIPTION_TEXT = 'Event Log:'

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Walk through all subfolders
pdf_files = []
for root, dirs, files in os.walk(PDF_FOLDER):
    for f in files:
        if f.lower().endswith('.pdf'):
            pdf_files.append(os.path.join(root, f))

print(f"Found {len(pdf_files)} PDF files in all subfolders.")

for pdf_path in pdf_files:
    pdf_file = os.path.basename(pdf_path)
    word_path = os.path.join(OUTPUT_FOLDER, pdf_file.replace('.pdf', '.docx'))

    print(f"Processing {pdf_file}...")

    # Convert PDF to Word
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    cv.close()

    doc = Document(word_path)

    # Add logo at the top, centered and bigger
    if os.path.exists(LOGO_PATH):
        pic_paragraph = doc.add_paragraph()
        run = pic_paragraph.add_run()
        run.add_picture(LOGO_PATH, width=Inches(3.5))
        pic_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Move picture paragraph to the start
        body = doc._body._element
        body.remove(pic_paragraph._element)
        body.insert(0, pic_paragraph._element)

    # Add bottom description in bold
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(DESCRIPTION_TEXT)
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(word_path)

print("All PDFs converted successfully with logo centered at the top!")
